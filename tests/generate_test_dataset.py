#!/usr/bin/env python3
"""Generate a comprehensive test dataset ZIP for the AI Grading System.

Creates 15 simulated student submissions covering every edge case:
- Multiple code languages (Python, Java, C++, JS)
- PDF (text + scanned/image-only)
- DOCX with tables
- Jupyter notebooks
- Images (PNG, JPG)
- Nested directory structures
- Empty files, unicode, unsupported formats, macOS junk
"""
import json
import os
import shutil
import zipfile
from pathlib import Path

# ---------------------------------------------------------------------------
# Config
# ---------------------------------------------------------------------------
SCRIPT_DIR = Path(__file__).parent
OUTPUT_ZIP = SCRIPT_DIR / "test_comprehensive_dataset.zip"
TEMP_DIR = SCRIPT_DIR / "_temp_dataset"

ASSIGNMENT_TITLE = "CS301 - Advanced Programming: BST & Analysis"
ASSIGNMENT_DESC = """Students must complete the following:
Problem 1: Implement a Binary Search Tree with insert, delete, search, and traversal methods.
Problem 2: Write a complexity analysis report covering time and space complexity.
Problem 3: Provide test cases with screenshots or evidence of output.
Submissions may include code (.py, .java, .cpp), documents (PDF/DOCX), images, or Jupyter notebooks."""

RUBRIC = """Correctness: 30 points
Code Quality: 20 points
Documentation: 20 points
Testing: 15 points
Analysis: 15 points"""


# ---------------------------------------------------------------------------
# Content generators
# ---------------------------------------------------------------------------

def _py_bst_good():
    return '''"""Binary Search Tree implementation — complete solution."""

class Node:
    """A node in the BST."""
    def __init__(self, key, value=None):
        self.key = key
        self.value = value
        self.left = None
        self.right = None

class BinarySearchTree:
    """BST with insert, delete, search, and traversal."""

    def __init__(self):
        self.root = None
        self.size = 0

    def insert(self, key, value=None):
        """Insert a key-value pair into the BST."""
        if self.root is None:
            self.root = Node(key, value)
        else:
            self._insert_recursive(self.root, key, value)
        self.size += 1

    def _insert_recursive(self, node, key, value):
        if key < node.key:
            if node.left is None:
                node.left = Node(key, value)
            else:
                self._insert_recursive(node.left, key, value)
        elif key > node.key:
            if node.right is None:
                node.right = Node(key, value)
            else:
                self._insert_recursive(node.right, key, value)
        else:
            node.value = value  # Update existing key

    def search(self, key):
        """Search for a key. Returns value or None."""
        return self._search_recursive(self.root, key)

    def _search_recursive(self, node, key):
        if node is None:
            return None
        if key == node.key:
            return node.value
        elif key < node.key:
            return self._search_recursive(node.left, key)
        else:
            return self._search_recursive(node.right, key)

    def delete(self, key):
        """Delete a node by key."""
        self.root = self._delete_recursive(self.root, key)

    def _delete_recursive(self, node, key):
        if node is None:
            return None
        if key < node.key:
            node.left = self._delete_recursive(node.left, key)
        elif key > node.key:
            node.right = self._delete_recursive(node.right, key)
        else:
            if node.left is None:
                self.size -= 1
                return node.right
            elif node.right is None:
                self.size -= 1
                return node.left
            successor = self._min_node(node.right)
            node.key = successor.key
            node.value = successor.value
            node.right = self._delete_recursive(node.right, successor.key)
        return node

    def _min_node(self, node):
        while node.left:
            node = node.left
        return node

    def inorder(self):
        """In-order traversal returns sorted keys."""
        result = []
        self._inorder_recursive(self.root, result)
        return result

    def _inorder_recursive(self, node, result):
        if node:
            self._inorder_recursive(node.left, result)
            result.append(node.key)
            self._inorder_recursive(node.right, result)

if __name__ == "__main__":
    bst = BinarySearchTree()
    for v in [50, 30, 70, 20, 40, 60, 80]:
        bst.insert(v)
    print("In-order:", bst.inorder())
    print("Search 40:", bst.search(40))
    bst.delete(30)
    print("After delete 30:", bst.inorder())
'''


def _py_bst_empty():
    return "# TODO: Implement BST\n"


def _py_bst_partial():
    return '''"""Partial BST — insert only, no delete."""

class BST:
    def __init__(self):
        self.root = None

    def insert(self, val):
        if not self.root:
            self.root = {"val": val, "left": None, "right": None}
        else:
            self._ins(self.root, val)

    def _ins(self, n, val):
        if val < n["val"]:
            if n["left"] is None:
                n["left"] = {"val": val, "left": None, "right": None}
            else:
                self._ins(n["left"], val)
        else:
            if n["right"] is None:
                n["right"] = {"val": val, "left": None, "right": None}
            else:
                self._ins(n["right"], val)
'''


def _py_tests():
    return '''"""Test cases for BST implementation."""
import unittest

class TestBST(unittest.TestCase):
    def test_insert_and_search(self):
        from q1_bst import BinarySearchTree
        bst = BinarySearchTree()
        bst.insert(10)
        bst.insert(5)
        bst.insert(15)
        self.assertEqual(bst.search(10), None)  # value is None
        self.assertIsNone(bst.search(99))

    def test_delete(self):
        from q1_bst import BinarySearchTree
        bst = BinarySearchTree()
        for v in [50, 30, 70]:
            bst.insert(v)
        bst.delete(30)
        self.assertEqual(bst.inorder(), [50, 70])

    def test_inorder(self):
        from q1_bst import BinarySearchTree
        bst = BinarySearchTree()
        for v in [5, 3, 7, 1, 4]:
            bst.insert(v)
        self.assertEqual(bst.inorder(), [1, 3, 4, 5, 7])

    def test_empty_tree(self):
        from q1_bst import BinarySearchTree
        bst = BinarySearchTree()
        self.assertIsNone(bst.search(1))
        self.assertEqual(bst.inorder(), [])

if __name__ == "__main__":
    unittest.main()
'''


def _java_bst():
    return '''// BinarySearchTree.java — Java implementation
public class BinarySearchTree {
    private class Node {
        int key;
        Node left, right;
        Node(int k) { key = k; left = right = null; }
    }

    private Node root;

    public void insert(int key) {
        root = insertRec(root, key);
    }

    private Node insertRec(Node root, int key) {
        if (root == null) return new Node(key);
        if (key < root.key) root.left = insertRec(root.left, key);
        else if (key > root.key) root.right = insertRec(root.right, key);
        return root;
    }

    public boolean search(int key) {
        return searchRec(root, key);
    }

    private boolean searchRec(Node root, int key) {
        if (root == null) return false;
        if (key == root.key) return true;
        return key < root.key ? searchRec(root.left, key) : searchRec(root.right, key);
    }

    public void delete(int key) {
        root = deleteRec(root, key);
    }

    private Node deleteRec(Node root, int key) {
        if (root == null) return null;
        if (key < root.key) root.left = deleteRec(root.left, key);
        else if (key > root.key) root.right = deleteRec(root.right, key);
        else {
            if (root.left == null) return root.right;
            if (root.right == null) return root.left;
            Node succ = minValue(root.right);
            root.key = succ.key;
            root.right = deleteRec(root.right, succ.key);
        }
        return root;
    }

    private Node minValue(Node node) {
        while (node.left != null) node = node.left;
        return node;
    }

    public void inorder() { inorderRec(root); System.out.println(); }
    private void inorderRec(Node root) {
        if (root != null) {
            inorderRec(root.left);
            System.out.print(root.key + " ");
            inorderRec(root.right);
        }
    }

    public static void main(String[] args) {
        BinarySearchTree bst = new BinarySearchTree();
        int[] keys = {50, 30, 70, 20, 40};
        for (int k : keys) bst.insert(k);
        System.out.print("Inorder: "); bst.inorder();
        System.out.println("Search 40: " + bst.search(40));
        bst.delete(30);
        System.out.print("After delete: "); bst.inorder();
    }
}
'''


def _java_test():
    return '''// TestBST.java
public class TestBST {
    public static void main(String[] args) {
        BinarySearchTree bst = new BinarySearchTree();
        bst.insert(10);
        bst.insert(5);
        bst.insert(15);
        assert bst.search(10) : "Search for 10 failed";
        assert !bst.search(99) : "Search for 99 should fail";
        bst.delete(10);
        assert !bst.search(10) : "10 should be deleted";
        System.out.println("All tests passed!");
    }
}
'''


def _cpp_bst():
    return '''// bst.cpp — C++ BST implementation
#include "bst.h"
#include <iostream>

BST::BST() : root(nullptr), sz(0) {}

BST::Node* BST::insert(Node* node, int key) {
    if (!node) { sz++; return new Node{key, nullptr, nullptr}; }
    if (key < node->key) node->left = insert(node->left, key);
    else if (key > node->key) node->right = insert(node->right, key);
    return node;
}

void BST::insert(int key) { root = insert(root, key); }

bool BST::search(int key) const {
    Node* cur = root;
    while (cur) {
        if (key == cur->key) return true;
        cur = key < cur->key ? cur->left : cur->right;
    }
    return false;
}

BST::Node* BST::findMin(Node* node) {
    while (node->left) node = node->left;
    return node;
}

BST::Node* BST::remove(Node* node, int key) {
    if (!node) return nullptr;
    if (key < node->key) node->left = remove(node->left, key);
    else if (key > node->key) node->right = remove(node->right, key);
    else {
        if (!node->left) { Node* r = node->right; delete node; sz--; return r; }
        if (!node->right) { Node* l = node->left; delete node; sz--; return l; }
        Node* succ = findMin(node->right);
        node->key = succ->key;
        node->right = remove(node->right, succ->key);
    }
    return node;
}

void BST::remove(int key) { root = remove(root, key); }

void BST::inorder(Node* node) const {
    if (node) { inorder(node->left); std::cout << node->key << " "; inorder(node->right); }
}

void BST::printInorder() const { inorder(root); std::cout << std::endl; }
int BST::size() const { return sz; }
'''


def _cpp_header():
    return '''// bst.h — BST header
#ifndef BST_H
#define BST_H

class BST {
    struct Node { int key; Node* left; Node* right; };
    Node* root;
    int sz;
    Node* insert(Node* node, int key);
    Node* remove(Node* node, int key);
    Node* findMin(Node* node);
    void inorder(Node* node) const;
public:
    BST();
    void insert(int key);
    bool search(int key) const;
    void remove(int key);
    void printInorder() const;
    int size() const;
};

#endif
'''


def _analysis_md():
    return '''# Complexity Analysis — Binary Search Tree

## Time Complexity

| Operation | Average Case | Worst Case |
|-----------|-------------|------------|
| Insert    | O(log n)    | O(n)       |
| Search    | O(log n)    | O(n)       |
| Delete    | O(log n)    | O(n)       |
| Traversal | O(n)        | O(n)       |

### Discussion
The average case assumes a balanced tree where height h = O(log n).
The worst case occurs with a degenerate (skewed) tree where all nodes
form a single chain, effectively becoming a linked list with h = O(n).

## Space Complexity
- **Tree storage:** O(n) for n nodes
- **Recursive operations:** O(h) stack space where h is the tree height
- **In-order traversal (iterative):** O(h) with explicit stack

## Balancing Considerations
Self-balancing variants like AVL trees or Red-Black trees guarantee
O(log n) worst-case for all operations by maintaining balance invariants.
'''


def _analysis_txt():
    return '''BST Complexity Analysis
=======================
Insert: O(log n) average, O(n) worst
Search: O(log n) average, O(n) worst  
Delete: O(log n) average, O(n) worst
Space:  O(n) overall, O(h) for recursion stack
The worst case happens when the tree degenerates into a linked list.
'''


def _unicode_py():
    return '''# -*- coding: utf-8 -*-
"""BST implementation with unicode comments — Ağaç Yapısı (Tree Structure)."""

class Düğüm:
    """Ağaç düğümü — Tree node."""
    def __init__(self, anahtar):
        self.anahtar = anahtar  # key
        self.sol = None         # left — 左
        self.sağ = None         # right — 右

class İkiliAramaAğacı:
    """İkili Arama Ağacı — Binary Search Tree — 二叉搜索树"""
    def __init__(self):
        self.kök = None  # root

    def ekle(self, anahtar):
        """Ekleme — Insert — 插入"""
        if self.kök is None:
            self.kök = Düğüm(anahtar)
        else:
            self._ekle(self.kök, anahtar)

    def _ekle(self, düğüm, anahtar):
        if anahtar < düğüm.anahtar:
            if düğüm.sol is None:
                düğüm.sol = Düğüm(anahtar)
            else:
                self._ekle(düğüm.sol, anahtar)
        else:
            if düğüm.sağ is None:
                düğüm.sağ = Düğüm(anahtar)
            else:
                self._ekle(düğüm.sağ, anahtar)

    def ara(self, anahtar):
        """Arama — Search — 搜索"""
        return self._ara(self.kök, anahtar)

    def _ara(self, düğüm, anahtar):
        if düğüm is None:
            return False
        if anahtar == düğüm.anahtar:
            return True
        elif anahtar < düğüm.anahtar:
            return self._ara(düğüm.sol, anahtar)
        else:
            return self._ara(düğüm.sağ, anahtar)

# Test: Çalıştırma — Execution — 実行
if __name__ == "__main__":
    ağaç = İkiliAramaAğacı()
    for değer in [50, 30, 70, 20, 40]:
        ağaç.ekle(değer)
    print(f"Arama 40: {ağaç.ara(40)}")  # True
    print(f"Arama 99: {ağaç.ara(99)}")  # False
'''


def _everything_py():
    return '''"""Single-file complete submission — all answers in one file."""

# ============ Problem 1: BST Implementation ============
class BST:
    def __init__(self): self.root = None
    def insert(self, v):
        if not self.root: self.root = {"v": v, "l": None, "r": None}
        else: self._ins(self.root, v)
    def _ins(self, n, v):
        if v < n["v"]:
            if not n["l"]: n["l"] = {"v": v, "l": None, "r": None}
            else: self._ins(n["l"], v)
        else:
            if not n["r"]: n["r"] = {"v": v, "l": None, "r": None}
            else: self._ins(n["r"], v)
    def search(self, v):
        return self._srch(self.root, v)
    def _srch(self, n, v):
        if not n: return False
        if v == n["v"]: return True
        return self._srch(n["l"], v) if v < n["v"] else self._srch(n["r"], v)

# ============ Problem 2: Analysis ============
# Insert: O(log n) avg, O(n) worst
# Search: O(log n) avg, O(n) worst
# Delete: not implemented
# Space: O(n)

# ============ Problem 3: Tests ============
if __name__ == "__main__":
    t = BST()
    for x in [5, 3, 7, 1, 4, 6, 8]:
        t.insert(x)
    assert t.search(4), "FAIL: 4 should exist"
    assert not t.search(99), "FAIL: 99 should not exist"
    print("All tests passed!")
'''


# ---------------------------------------------------------------------------
# Binary file generators (PDF, DOCX, Notebook, Image)
# ---------------------------------------------------------------------------

def _make_text_pdf(path: Path, title: str, body: str):
    """Create a text-based PDF using PyMuPDF."""
    import fitz
    doc = fitz.open()
    page = doc.new_page(width=595, height=842)  # A4
    # Title
    page.insert_text((50, 60), title, fontsize=18, fontname="helv", color=(0, 0, 0.6))
    # Body text — split into lines to fit page
    y = 100
    for line in body.split("\n"):
        if y > 790:
            page = doc.new_page(width=595, height=842)
            y = 50
        page.insert_text((50, y), line, fontsize=10, fontname="helv")
        y += 14
    doc.save(str(path))
    doc.close()


def _make_scanned_pdf(path: Path):
    """Create an image-only PDF (simulating handwritten/scanned) — minimal text extraction."""
    import fitz
    from PIL import Image, ImageDraw, ImageFont
    import io

    # Create an image that looks like handwriting
    img = Image.new("RGB", (595, 842), (255, 253, 240))  # cream paper
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Noteworthy.ttc", 18)
    except Exception:
        font = ImageFont.load_default()

    lines = [
        "BST Complexity Analysis",
        "",
        "Insert: O(log n) average case",
        "        O(n) worst case (skewed)",
        "",
        "Search: O(log n) average",
        "        O(n) worst case",
        "",
        "Delete: O(log n) average",
        "        O(n) worst case",
        "",
        "Space: O(n) for storage",
        "       O(h) for recursion stack",
        "",
        "Note: self-balancing trees like",
        "AVL guarantee O(log n) worst case",
    ]
    y = 80
    for line in lines:
        # Slight random offset for handwritten feel
        draw.text((60, y), line, fill=(30, 30, 80), font=font)
        y += 32

    # Add some "doodles" — a simple tree diagram
    draw.ellipse([250, 550, 290, 590], outline=(0, 0, 100), width=2)
    draw.text((260, 558), "50", fill=(0, 0, 100), font=font)
    draw.line([270, 590, 220, 640], fill=(0, 0, 100), width=2)
    draw.line([270, 590, 320, 640], fill=(0, 0, 100), width=2)
    draw.ellipse([200, 640, 240, 680], outline=(0, 0, 100), width=2)
    draw.text((210, 648), "30", fill=(0, 0, 100), font=font)
    draw.ellipse([300, 640, 340, 680], outline=(0, 0, 100), width=2)
    draw.text((310, 648), "70", fill=(0, 0, 100), font=font)

    # Convert to PDF via fitz
    img_bytes = io.BytesIO()
    img.save(img_bytes, format="PNG")
    img_bytes.seek(0)

    doc = fitz.open()
    page = doc.new_page(width=595, height=842)
    page.insert_image(fitz.Rect(0, 0, 595, 842), stream=img_bytes.read())
    doc.save(str(path))
    doc.close()


def _make_docx(path: Path, title: str, body_paragraphs: list, table_data: list = None):
    """Create a DOCX with paragraphs and optional tables."""
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for para in body_paragraphs:
        doc.add_paragraph(para)

    if table_data:
        doc.add_heading("Complexity Table", level=2)
        table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
        table.style = "Table Grid"
        for i, row in enumerate(table_data):
            for j, cell in enumerate(row):
                table.rows[i].cells[j].text = cell

    doc.save(str(path))


def _make_notebook(path: Path, cells: list):
    """Create a Jupyter notebook. cells = list of (type, source) tuples."""
    import nbformat
    nb = nbformat.v4.new_notebook()
    for ctype, source in cells:
        if ctype == "markdown":
            nb.cells.append(nbformat.v4.new_markdown_cell(source))
        else:
            nb.cells.append(nbformat.v4.new_code_cell(source))
    with open(path, "w") as f:
        nbformat.write(nb, f)


def _make_image(path: Path, label: str, w=640, h=480, bg=(255, 255, 255)):
    """Create a simple diagram/screenshot image."""
    from PIL import Image, ImageDraw, ImageFont
    img = Image.new("RGB", (w, h), bg)
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 16)
        font_sm = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 12)
    except Exception:
        font = ImageFont.load_default()
        font_sm = font

    # Header bar
    draw.rectangle([0, 0, w, 40], fill=(50, 50, 120))
    draw.text((10, 10), label, fill=(255, 255, 255), font=font)

    # Tree diagram
    cx, cy = w // 2, 120
    nodes = [(cx, cy, "50"), (cx - 100, cy + 80, "30"), (cx + 100, cy + 80, "70"),
             (cx - 150, cy + 160, "20"), (cx - 50, cy + 160, "40")]
    edges = [(0, 1), (0, 2), (1, 3), (1, 4)]
    for a, b in edges:
        draw.line([nodes[a][0], nodes[a][1], nodes[b][0], nodes[b][1]], fill=(100, 100, 100), width=2)
    for x, y, txt in nodes:
        draw.ellipse([x - 20, y - 20, x + 20, y + 20], fill=(100, 149, 237), outline=(50, 50, 120), width=2)
        draw.text((x - 8, y - 8), txt, fill=(255, 255, 255), font=font_sm)

    # Console output area
    draw.rectangle([20, cy + 220, w - 20, h - 20], fill=(30, 30, 50))
    console_lines = [
        ">>> bst.inorder()",
        "[20, 30, 40, 50, 70]",
        ">>> bst.search(40)",
        "True",
        ">>> bst.delete(30)",
        ">>> bst.inorder()",
        "[20, 40, 50, 70]",
    ]
    ty = cy + 235
    for line in console_lines:
        color = (0, 255, 100) if not line.startswith(">>>") else (200, 200, 200)
        draw.text((30, ty), line, fill=color, font=font_sm)
        ty += 18

    ext = path.suffix.lower()
    fmt = "JPEG" if ext in (".jpg", ".jpeg") else "PNG"
    img.save(str(path), format=fmt)


# ---------------------------------------------------------------------------
# Student builders
# ---------------------------------------------------------------------------

def _build_alice_perfect(base: Path):
    """Perfect Python submission — ideal case."""
    d = base / "alice_perfect"
    d.mkdir(parents=True)
    (d / "q1_bst.py").write_text(_py_bst_good(), encoding="utf-8")
    (d / "q2_analysis.md").write_text(_analysis_md(), encoding="utf-8")
    (d / "q3_tests.py").write_text(_py_tests(), encoding="utf-8")


def _build_bob_java(base: Path):
    """Java submission — non-Python language."""
    d = base / "bob_java"
    d.mkdir(parents=True)
    (d / "BinarySearchTree.java").write_text(_java_bst(), encoding="utf-8")
    (d / "analysis.txt").write_text(_analysis_txt(), encoding="utf-8")
    (d / "TestBST.java").write_text(_java_test(), encoding="utf-8")


def _build_carol_cpp(base: Path):
    """C++ with header + DOCX report."""
    d = base / "carol_cpp"
    d.mkdir(parents=True)
    (d / "bst.cpp").write_text(_cpp_bst(), encoding="utf-8")
    (d / "bst.h").write_text(_cpp_header(), encoding="utf-8")
    _make_docx(
        d / "report.docx",
        "BST Complexity Analysis Report",
        [
            "This report analyses the time and space complexity of BST operations.",
            "Insert, search, and delete all run in O(log n) average case.",
            "The worst case for all operations is O(n) for a degenerate tree.",
            "Space complexity is O(n) for storage plus O(h) for recursion stack.",
        ],
        table_data=[
            ["Operation", "Average", "Worst"],
            ["Insert", "O(log n)", "O(n)"],
            ["Search", "O(log n)", "O(n)"],
            ["Delete", "O(log n)", "O(n)"],
            ["Traversal", "O(n)", "O(n)"],
        ]
    )


def _build_dan_docx(base: Path):
    """DOCX with tables and formatted text."""
    d = base / "dan_docx"
    d.mkdir(parents=True)
    (d / "solution.py").write_text(_py_bst_good(), encoding="utf-8")
    _make_docx(
        d / "Q2_report.docx",
        "Problem 2: Algorithm Analysis",
        [
            "Binary Search Trees maintain the invariant that left < root < right.",
            "This ordering property enables efficient searching in O(log n) time.",
            "Deletion is the most complex operation with three cases to handle:",
            "1. Leaf node: simply remove",
            "2. One child: replace with child",
            "3. Two children: replace with inorder successor",
        ],
        table_data=[
            ["Operation", "Best", "Average", "Worst", "Space"],
            ["Insert", "O(1)", "O(log n)", "O(n)", "O(1)"],
            ["Search", "O(1)", "O(log n)", "O(n)", "O(1)"],
            ["Delete", "O(1)", "O(log n)", "O(n)", "O(1)"],
            ["Traversal", "O(n)", "O(n)", "O(n)", "O(n)"],
        ]
    )


def _build_eve_pdf_text(base: Path):
    """PDF with extractable text — tests _parse_pdf text path."""
    d = base / "eve_pdf_text"
    d.mkdir(parents=True)
    (d / "bst_implementation.py").write_text(_py_bst_good(), encoding="utf-8")
    _make_text_pdf(
        d / "analysis_report.pdf",
        "BST Complexity Analysis",
        """Time Complexity Analysis
========================

Insert Operation:
- Average case: O(log n) — the tree is roughly balanced
- Worst case: O(n) — all elements inserted in sorted order

Search Operation:
- Average case: O(log n) — binary search on tree structure
- Worst case: O(n) — degenerate tree

Delete Operation:
- Average case: O(log n)
- Worst case: O(n)
- Requires handling three cases: leaf, one child, two children

Traversal:
- In-order: O(n) — visits every node exactly once
- Produces sorted output

Space Complexity:
- Storage: O(n) for n nodes
- Recursion: O(h) stack frames, where h is tree height
- h = O(log n) balanced, O(n) worst case

Self-Balancing Trees:
AVL and Red-Black trees guarantee O(log n) height through
rotations after insert/delete, ensuring worst-case O(log n)
for all operations at the cost of more complex code."""
    )


def _build_frank_scanned(base: Path):
    """Scanned/image-only PDF — tests _pdf_to_images fallback."""
    d = base / "frank_pdf_scanned"
    d.mkdir(parents=True)
    (d / "code.py").write_text(_py_bst_partial(), encoding="utf-8")
    _make_scanned_pdf(d / "handwritten_analysis.pdf")


def _build_grace_notebook(base: Path):
    """Jupyter notebook with code + markdown cells."""
    d = base / "grace_notebook"
    d.mkdir(parents=True)
    _make_notebook(d / "assignment.ipynb", [
        ("markdown", "# Problem 1: Binary Search Tree\nComplete BST implementation with insert, delete, and search."),
        ("code", '''class Node:
    def __init__(self, key):
        self.key = key
        self.left = None
        self.right = None

class BST:
    def __init__(self):
        self.root = None

    def insert(self, key):
        if not self.root:
            self.root = Node(key)
        else:
            self._insert(self.root, key)

    def _insert(self, node, key):
        if key < node.key:
            if node.left is None: node.left = Node(key)
            else: self._insert(node.left, key)
        else:
            if node.right is None: node.right = Node(key)
            else: self._insert(node.right, key)

    def search(self, key):
        return self._search(self.root, key)

    def _search(self, node, key):
        if not node: return False
        if key == node.key: return True
        return self._search(node.left, key) if key < node.key else self._search(node.right, key)

    def delete(self, key):
        self.root = self._delete(self.root, key)

    def _delete(self, node, key):
        if not node: return None
        if key < node.key: node.left = self._delete(node.left, key)
        elif key > node.key: node.right = self._delete(node.right, key)
        else:
            if not node.left: return node.right
            if not node.right: return node.left
            succ = node.right
            while succ.left: succ = succ.left
            node.key = succ.key
            node.right = self._delete(node.right, succ.key)
        return node

    def inorder(self):
        result = []
        self._inorder(self.root, result)
        return result

    def _inorder(self, node, result):
        if node:
            self._inorder(node.left, result)
            result.append(node.key)
            self._inorder(node.right, result)'''),
        ("markdown", "## Testing"),
        ("code", '''bst = BST()
for v in [50, 30, 70, 20, 40, 60, 80]:
    bst.insert(v)
print("Inorder:", bst.inorder())
print("Search 40:", bst.search(40))
bst.delete(30)
print("After delete:", bst.inorder())'''),
        ("markdown", "## Problem 2: Complexity Analysis\n| Op | Avg | Worst |\n|---|---|---|\n| Insert | O(log n) | O(n) |\n| Search | O(log n) | O(n) |\n| Delete | O(log n) | O(n) |"),
        ("markdown", "## Conclusion\nThe BST provides efficient average-case performance but degrades to O(n) for skewed inputs. Self-balancing variants address this limitation."),
    ])


def _build_henry_images(base: Path):
    """Screenshots and images — tests multimodal vision pipeline."""
    d = base / "henry_images"
    d.mkdir(parents=True)
    (d / "solution.py").write_text(_py_bst_good(), encoding="utf-8")
    _make_image(d / "test_output.png", "BST Test Output — test_bst.py")
    _make_image(d / "analysis_diagram.jpg", "BST Diagram & Analysis", w=800, h=600, bg=(245, 245, 255))


def _build_iris_nested(base: Path):
    """Deeply nested directory structure."""
    d = base / "iris_nested"
    # Nested code
    src = d / "src" / "main" / "java" / "bst"
    src.mkdir(parents=True)
    (src / "BST.java").write_text(_java_bst(), encoding="utf-8")
    # Nested docs
    docs = d / "docs"
    docs.mkdir(parents=True)
    _make_text_pdf(docs / "report.pdf", "BST Analysis", "Insert: O(log n)\nSearch: O(log n)\nDelete: O(log n)\nSpace: O(n)\n")
    # Nested tests
    tests = d / "tests"
    tests.mkdir(parents=True)
    (tests / "test_bst.py").write_text(_py_tests(), encoding="utf-8")


def _build_jake_flat(base: Path):
    """Single file — everything in one."""
    d = base / "jake_flat"
    d.mkdir(parents=True)
    (d / "everything.py").write_text(_everything_py(), encoding="utf-8")


def _build_karen_junk(base: Path):
    """Valid files mixed with macOS junk that should be filtered."""
    d = base / "karen_macos_junk"
    d.mkdir(parents=True)
    (d / "solution.py").write_text(_py_bst_good(), encoding="utf-8")
    (d / "analysis.md").write_text(_analysis_md(), encoding="utf-8")
    # macOS junk
    macos = d / "__MACOSX"
    macos.mkdir()
    (macos / "._solution.py").write_text("mac resource fork garbage", encoding="utf-8")
    (d / ".DS_Store").write_bytes(b"\x00\x00\x00\x01Bud1" + b"\x00" * 100)
    pycache = d / "__pycache__"
    pycache.mkdir()
    (pycache / "solution.cpython-311.pyc").write_bytes(b"\x00" * 50)


def _build_leo_empty(base: Path):
    """Empty/zero-byte files — edge case for all parsers."""
    d = base / "leo_empty"
    d.mkdir(parents=True)
    (d / "q1.py").write_text("", encoding="utf-8")
    (d / "q2.txt").write_text("", encoding="utf-8")
    (d / "q3.py").write_text("", encoding="utf-8")


def _build_mia_unicode(base: Path):
    """Unicode filenames and content."""
    d = base / "mia_unicode"
    d.mkdir(parents=True)
    (d / "solution_v2.py").write_text(_unicode_py(), encoding="utf-8")
    (d / "analysis_notes.txt").write_text(
        "分析：BST的时间复杂度\n"
        "• 插入: O(log n) 平均, O(n) 最坏\n"
        "• 搜索: O(log n) 平均, O(n) 最坏\n"
        "• 删除: O(log n) 平均, O(n) 最坏\n"
        "• 空间: O(n)\n"
        "\nТакже: сбалансированные деревья (AVL, Red-Black)\n"
        "обеспечивают O(log n) в худшем случае.\n",
        encoding="utf-8"
    )


def _build_nick_mixed(base: Path):
    """Every supported format at once."""
    d = base / "nick_mixed"
    d.mkdir(parents=True)
    (d / "q1_bst.py").write_text(_py_bst_good(), encoding="utf-8")
    _make_text_pdf(d / "q2_report.pdf", "BST Report",
                   "Insert O(log n), Search O(log n), Delete O(log n)")
    _make_docx(d / "extra_notes.docx", "Additional Notes",
               ["BST is a fundamental data structure.", "Self-balancing variants improve worst-case."])
    _make_notebook(d / "experiments.ipynb", [
        ("markdown", "# BST Experiments"),
        ("code", "print('Testing BST performance...')"),
    ])
    _make_image(d / "output_screenshot.png", "BST Test Results")
    (d / "readme.txt").write_text("Nick's submission — all formats included.", encoding="utf-8")
    _make_image(d / "diagram.jpg", "BST Structure Diagram", bg=(240, 248, 255))


def _build_olivia_unsupported(base: Path):
    """Mix of valid and unsupported file types."""
    d = base / "olivia_unsupported"
    d.mkdir(parents=True)
    (d / "solution.py").write_text(_py_bst_partial(), encoding="utf-8")
    # Unsupported formats
    (d / "data.csv").write_text("key,left,right\n50,30,70\n30,20,40\n", encoding="utf-8")
    (d / "config.xml").write_text('<?xml version="1.0"?>\n<bst><node key="50"/></bst>\n', encoding="utf-8")
    (d / "notes.rtf").write_text(r"{\rtf1 BST notes}", encoding="utf-8")
    (d / "model.pkl").write_bytes(b"\x80\x04\x95" + b"\x00" * 50)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def build_dataset():
    """Build the complete test dataset and package as ZIP."""
    print(f"Building comprehensive test dataset...")

    # Clean up
    if TEMP_DIR.exists():
        shutil.rmtree(TEMP_DIR)
    TEMP_DIR.mkdir(parents=True)

    builders = [
        ("alice_perfect",       _build_alice_perfect),
        ("bob_java",            _build_bob_java),
        ("carol_cpp",           _build_carol_cpp),
        ("dan_docx",            _build_dan_docx),
        ("eve_pdf_text",        _build_eve_pdf_text),
        ("frank_pdf_scanned",   _build_frank_scanned),
        ("grace_notebook",      _build_grace_notebook),
        ("henry_images",        _build_henry_images),
        ("iris_nested",         _build_iris_nested),
        ("jake_flat",           _build_jake_flat),
        ("karen_macos_junk",    _build_karen_junk),
        ("leo_empty",           _build_leo_empty),
        ("mia_unicode",         _build_mia_unicode),
        ("nick_mixed",          _build_nick_mixed),
        ("olivia_unsupported",  _build_olivia_unsupported),
    ]

    for name, builder in builders:
        try:
            builder(TEMP_DIR)
            # Count files
            count = sum(1 for _ in (TEMP_DIR / name).rglob("*") if _.is_file())
            print(f"  ✅ {name:25s} ({count} files)")
        except Exception as e:
            print(f"  ❌ {name:25s} FAILED: {e}")

    # Package as ZIP
    print(f"\nPackaging into {OUTPUT_ZIP.name}...")
    if OUTPUT_ZIP.exists():
        OUTPUT_ZIP.unlink()
    with zipfile.ZipFile(OUTPUT_ZIP, "w", zipfile.ZIP_DEFLATED) as zf:
        for file in sorted(TEMP_DIR.rglob("*")):
            if file.is_file():
                arcname = file.relative_to(TEMP_DIR)
                zf.write(file, arcname)

    # Stats
    with zipfile.ZipFile(OUTPUT_ZIP, "r") as zf:
        total_files = len(zf.namelist())

    size_mb = OUTPUT_ZIP.stat().st_size / (1024 * 1024)
    print(f"\n{'='*50}")
    print(f"Dataset generated: {OUTPUT_ZIP}")
    print(f"Total files: {total_files}")
    print(f"ZIP size: {size_mb:.2f} MB")
    print(f"{'='*50}")

    # Print assignment info for session creation
    print(f"\n📋 Assignment Info for Session Creation:")
    print(f"Title: {ASSIGNMENT_TITLE}")
    print(f"Max Score: 100")
    print(f"\nDescription:\n{ASSIGNMENT_DESC}")
    print(f"\nRubric:\n{RUBRIC}")

    # Cleanup temp
    shutil.rmtree(TEMP_DIR)
    print(f"\n✅ Temp files cleaned up. ZIP ready for upload!")


if __name__ == "__main__":
    build_dataset()
