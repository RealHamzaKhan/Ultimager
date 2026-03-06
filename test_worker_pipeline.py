#!/usr/bin/env python3
"""Regression tests for provider routing, failover, and large-image stability."""

from __future__ import annotations

import asyncio
import base64
import json
from dataclasses import dataclass, field
from pathlib import Path

import pytest

from app.services import ai_grader_fixed as grader
from app.services.file_parser_fixed import process_student_submission


@dataclass
class _MockContent:
    filename: str
    file_type: str = "pdf"
    text_content: str = ""
    images: list[dict] = field(default_factory=list)


def _img_b64(seed: str) -> str:
    return base64.b64encode(f"img:{seed}".encode("utf-8")).decode("utf-8")


def _make_image(file_idx: int, page_idx: int, kind: str) -> dict:
    is_focus = kind == "focus"
    return {
        "page": page_idx,
        "description": f"Page {page_idx} {kind}",
        "region_type": "embedded_focus" if is_focus else "full_page",
        "media_type": "image/png",
        "size_bytes": 1024 + file_idx + page_idx,
        "base64": _img_b64(f"{file_idx}-{page_idx}-{kind}"),
    }


def test_more_than_50_images_coverage_is_preserved():
    # 8 files * 8 pages * 2 views = 128 images
    files = []
    for fi in range(1, 9):
        images = []
        for pi in range(1, 9):
            images.append(_make_image(fi, pi, "focus"))
            images.append(_make_image(fi, pi, "full"))
        files.append(_MockContent(filename=f"student_{fi}.pdf", images=images))

    selected = grader._collect_selected_images(files, max_images=200)
    assert len(selected) == 128
    assert any(img["filename"] == "student_8.pdf" and img["page"] == 8 for img in selected)
    assert len({img["image_id"] for img in selected}) == 128

    pre = grader._select_images_for_preanalysis(selected, max_images=200)
    assert len(pre) == 128


def test_mixed_pdf_docx_png_ingestion_with_handwritten_signals(tmp_path: Path):
    reportlab = pytest.importorskip("reportlab")
    pytest.importorskip("docx")
    pil = pytest.importorskip("PIL.Image")
    _ = reportlab, pil

    student_dir = tmp_path / "student_mixed"
    student_dir.mkdir(parents=True, exist_ok=True)

    # PNG (handwritten-like label)
    from PIL import Image, ImageDraw

    img_path = student_dir / "diagram.png"
    img = Image.new("RGB", (600, 320), color=(255, 255, 255))
    draw = ImageDraw.Draw(img)
    draw.text((40, 120), "handwritten flow sketch", fill=(15, 15, 15))
    img.save(img_path)

    # DOCX
    from docx import Document

    doc = Document()
    doc.add_paragraph("Handwritten notes transcribed: BFS queue updates")
    doc.add_paragraph("Diagram explanation for balancing case.")
    doc_path = student_dir / "analysis.docx"
    doc.save(doc_path)

    # PDF
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas

    pdf_path = student_dir / "submission.pdf"
    c = canvas.Canvas(str(pdf_path), pagesize=letter)
    c.drawString(90, 730, "Page 1: handwritten derivation (simulated)")
    c.drawString(90, 700, "Graph edges and BFS traversal notes")
    c.showPage()
    c.drawString(90, 730, "Page 2: diagram annotations")
    c.save()

    extracted, ingestion = process_student_submission(student_dir, "student_mixed", 1)
    names = {e.filename for e in extracted}

    assert "diagram.png" in names
    assert "analysis.docx" in names
    assert "submission.pdf" in names
    assert ingestion.total_images >= 2
    assert len(ingestion.files_parsed) >= 3


def test_failover_chain_ollama_to_openrouter(monkeypatch):
    specs = [
        grader.ProviderSpec("ollama", "http://localhost:11434/v1", "", "m1", "m1", {}, False),
        grader.ProviderSpec("groq", "https://groq.example", "k", "m3", "m3", {}, True),
        grader.ProviderSpec("openrouter", "https://or.example", "k", "m4", "m4", {}, True),
    ]

    monkeypatch.setattr(
        grader,
        "_enabled_provider_candidates",
        lambda needs_vision: [(spec, spec.model_text) for spec in specs],
    )
    monkeypatch.setattr(grader, "_get_client", lambda spec: object())
    monkeypatch.setattr(grader, "_apply_provider_cooldown", lambda provider_name, error_type: None)

    call_order: list[str] = []

    def fake_chat(client, *, _spec, purpose, provider_name, model, messages, temperature, max_tokens, top_p=None, seed=None):
        _ = client, purpose, model, messages, temperature, max_tokens, top_p, seed
        call_order.append(provider_name)
        if provider_name == "ollama":
            raise RuntimeError("request timeout while waiting for provider")
        if provider_name == "groq":
            raise RuntimeError("model not supported on this endpoint")
        return grader._ResponseShim('{"ok": true}')

    monkeypatch.setattr(grader, "_chat_completion", fake_chat)

    response, meta = grader._chat_completion_with_failover(
        purpose="grade_student",
        needs_vision=False,
        messages=[{"role": "user", "content": "ping"}],
        temperature=0.0,
        max_tokens=32,
        seed=42,
    )

    assert call_order == ["ollama", "groq", "openrouter"]
    assert meta["provider"] == "openrouter"
    assert (response.choices[0].message.content or "").strip().startswith("{")


def test_consistency_across_regrade_runs(monkeypatch):
    async def no_wait():
        return None

    async def fake_vision_preanalysis(selected_images):
        entries = [{
            "image_id": img["image_id"],
            "summary": "Visible handwritten reasoning and diagram edges.",
            "transcription": "Queue: A,B,C ...",
            "substantive": True,
            "confidence": "high",
        } for img in selected_images]
        return (
            "Consolidated visual notes with image_id anchors.",
            {
                "enabled": True,
                "images_analyzed": len(selected_images),
                "chunks": 1,
                "batch_notes": [{
                    "batch_id": 1,
                    "image_ids": [img["image_id"] for img in selected_images],
                    "entries": entries,
                    "notes": "batch notes",
                }],
            },
        )

    def fake_chat_completion_with_failover(
        *,
        purpose,
        needs_vision,
        messages,
        temperature,
        max_tokens,
        top_p=None,
        seed=None,
        preferred_provider=None,
        allow_fallback=True,
    ):
        _ = purpose, needs_vision, messages, temperature, max_tokens, top_p, seed, preferred_provider, allow_fallback
        payload = {
            "rubric_breakdown": [
                {"criterion": "Algorithm Correctness", "score": 7, "max": 10, "justification": "Core logic works."},
                {"criterion": "Analysis Quality", "score": 8, "max": 10, "justification": "Reasoning present."},
            ],
            "total_score": 15,
            "overall_feedback": "Solid effort with clear evidence.",
            "strengths": ["Correct traversal implementation"],
            "weaknesses": ["Could improve edge-case discussion"],
            "suggestions_for_improvement": "Add explicit complexity analysis.",
            "confidence": "high",
            "confidence_reasoning": "Deterministic mock output.",
        }
        return (
            grader._ResponseShim(json.dumps(payload)),
            {"provider": "openrouter", "provider_key": "openrouter", "model": "mock-model", "attempts_before_success": [], "fallback_used": True},
        )

    monkeypatch.setattr(grader._rate_limiter, "acquire", no_wait)
    monkeypatch.setattr(grader, "_run_vision_preanalysis", fake_vision_preanalysis)
    monkeypatch.setattr(grader, "_chat_completion_with_failover", fake_chat_completion_with_failover)

    student_files = [
        _MockContent(
            filename="submission.pdf",
            file_type="pdf",
            text_content="BFS and UCS implementations with comments.",
            images=[_make_image(1, 1, "focus"), _make_image(1, 1, "full"), _make_image(1, 2, "focus")],
        )
    ]

    kwargs = dict(
        title="Graph Search Assignment",
        description="Implement BFS/UCS and explain tradeoffs.",
        rubric="Algorithm Correctness: 10 points\nAnalysis Quality: 10 points\nTotal: 20",
        max_score=20,
        student_files=student_files,
        questions=[],
    )

    first = asyncio.run(grader.grade_student(**kwargs))
    second = asyncio.run(grader.grade_student(**kwargs))

    assert first["grading_hash"] == second["grading_hash"]
    assert first["total_score"] == second["total_score"]
    assert first["rubric_breakdown"] == second["rubric_breakdown"]
    assert all(item.get("citations") for item in first.get("rubric_breakdown", []))
    assert first.get("transparency", {}).get("images_processed_in_batches", 0) >= 1


def test_relevance_gate_blocks_strong_off_topic_cases():
    relevance = {
        "is_relevant": False,
        "confidence": "high",
        "flags": ["off_topic", "wrong_assignment"],
        "reasoning": "Submission content is about an unrelated subject.",
    }

    gate = grader.evaluate_relevance_gate(relevance)
    assert gate["block_grading"] is True
    assert "off_topic" in gate["critical_flags"]
    assert "wrong_assignment" in gate["critical_flags"]

    result = grader.build_relevance_block_result(
        rubric="Algorithm Correctness: 8 points\nAnalysis Quality: 2 points\nTotal: 10",
        max_score=10,
        relevance=relevance,
        gate=gate,
    )
    assert result["total_score"] == 0.0
    assert result["letter_grade"] == "F"
    assert result.get("relevance_gate", {}).get("block_grading") is True
    assert len(result.get("rubric_breakdown", [])) == 2


def test_relevance_gate_allows_mixed_content_with_relevant_sections():
    relevance = {
        "is_relevant": False,
        "confidence": "high",
        "flags": ["off_topic", "wrong_assignment", "mixed_content"],
        "reasoning": "Some parts look unrelated, but assignment-specific work is present.",
        "has_relevant_sections": True,
        "assignment_signal": {
            "token_overlap": 7,
            "signal_ratio": 0.21,
            "files_with_overlap": 2,
            "files_scanned": 3,
            "matched_terms": ["bfs", "ucs", "manhattan"],
            "has_relevant_sections": True,
        },
    }

    gate = grader.evaluate_relevance_gate(relevance)
    assert gate["block_grading"] is False
    assert gate["review_required"] is True
    assert gate["mixed_content"] is True
    assert gate["has_relevant_sections"] is True


def test_citation_assignment_does_not_force_unrelated_image():
    evidence_map = [{
        "image_id": "img_0001",
        "filename": "24P-0547-Question#01.pdf",
        "page": 1,
        "description": "Campus navigation graph",
        "summary": "Graph nodes and BFS/UCS traversal notes",
        "transcription": "BFS queue updates and shortest path",
        "content_score": 42.0,
        "substantive": True,
        "confidence": "high",
    }]
    rubric_criteria = [{"criterion": "8-Puzzle Solver Implementation", "max": 2}]
    plan = grader._build_criterion_evidence_plan(rubric_criteria, text_content="", evidence_map=evidence_map)
    candidate_map = plan.get("candidate_map", {})

    key = "8-puzzle solver implementation"
    assert key in candidate_map
    assert candidate_map[key].get("image_ids", []) == []

    result = {
        "rubric_breakdown": [{
            "criterion": "8-Puzzle Solver Implementation",
            "score": 0,
            "max": 2,
            "justification": "No visible evidence.",
            "citations": [{"type": "image", "image_id": "img_0001"}],
        }]
    }
    stats = grader._attach_rubric_citations(result, evidence_map, candidate_map)
    citations = result["rubric_breakdown"][0].get("citations", [])

    assert stats["criteria_with_image_citation"] == 0
    assert all(not c.get("image_id") for c in citations)
    assert any(c.get("source") == "text_content" for c in citations)
